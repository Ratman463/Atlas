"""
Generate a Word (.docx) document from a Markdown-style outline produced by an
LLM. We keep the markup intentionally tiny so the model behaves:

    # 大标题           -> Heading 1
    ## 章节            -> Heading 2
    ### 小节           -> Heading 3
    - 项目 / * 项目   -> List Bullet
    1. 项目            -> List Number
    | a | b |          -> simple table
    > 引用             -> Quote
    其他                -> Paragraph (with **bold** / *italic* / `code`)

This avoids pulling in pandoc and keeps the whole project a single pip install.
"""
from __future__ import annotations

import io
import re
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)")
_CODE_RE = re.compile(r"`([^`]+?)`")


def _add_inline_runs(paragraph, text: str) -> None:
    """Render inline **bold**, *italic*, `code` markers into runs."""
    # Tokenise into (text, bold, italic, code) chunks.
    tokens: List[List] = [["", False, False, False]]

    def push(chunk_text, b=False, i=False, c=False):
        if not chunk_text:
            return
        tokens.append([chunk_text, b, i, c])
        tokens.append(["", False, False, False])

    def consume(pattern, flag_fn):
        # Process the most recent text token
        idx = len(tokens) - 1
        last = tokens[idx][0]
        if not last:
            return False
        m = pattern.search(last)
        if not m:
            return False
        before, inner, after = last[: m.start()], m.group(1), last[m.end():]
        tokens[idx][0] = before
        flag = flag_fn()
        push(inner, **({"c": True} if flag == "c" else
                       {"b": True} if flag == "b" else
                       {"i": True}))
        # Continue parsing the remainder in a fresh tail token
        tokens.append([after, False, False, False])
        return True

    # Walk the simple way: repeatedly scan the latest text token.
    progress = True
    while progress:
        progress = False
        # Try code first (so ** inside `..` stays literal)
        if consume(_CODE_RE, lambda: "c"):
            progress = True
        elif consume(_BOLD_RE, lambda: "b"):
            progress = True
        elif consume(_ITALIC_RE, lambda: "i"):
            progress = True

    for text, bold, italic, code in tokens:
        if not text:
            continue
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.color.rgb = RGBColor(0xA0, 0x2C, 0x2C)


def _add_table(doc: Document, header_cells: List[str], rows: List[List[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    table.style = "Light Grid Accent 1"
    try:
        table.autofit = True
    except Exception:
        pass

    for i, txt in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt.strip())
        run.bold = True

    for r, row in enumerate(rows, start=1):
        for c, txt in enumerate(row):
            if c >= len(header_cells):
                break
            cell = table.rows[r].cells[c]
            cell.text = ""
            cell.paragraphs[0].add_run(txt.strip())
    doc.add_paragraph("")  # spacing after table


def markdown_to_docx(title: str, markdown: str) -> bytes:
    """Convert Markdown-ish text to a .docx and return its bytes."""
    doc = Document()

    # Base style: clean, Apple-ish typography.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # Document title
    if title:
        h = doc.add_heading(level=0)
        run = h.add_run(title)
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph("")

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Skip blank lines
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("---") and len(line.strip()) == 3:
            # Horizontal rule -> empty paragraph with bottom border
            p = doc.add_paragraph("")
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:].strip())
            p.style = doc.styles["Intense Quote"] if "Intense Quote" in [s.name for s in doc.styles] else p.style
        elif re.match(r"^\s*[-*]\s+", line):
            item = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, item)
        elif re.match(r"^\s*\d+\.\s+", line):
            item = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, item)
        elif "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{2,}", lines[i + 1].strip().replace("|", "").replace(" ", "")):
            # Markdown table: header | separator | rows
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 1  # skip separator
            rows: List[List[str]] = []
            i += 1
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            _add_table(doc, header_cells, rows)
            continue
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line.strip())

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
